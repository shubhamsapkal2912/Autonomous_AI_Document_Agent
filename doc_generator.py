import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def safe_filename(name: str) -> str:
    name = name.lower().strip()
    name = re.sub(r"[^a-z0-9\s_-]", "", name)
    name = re.sub(r"[\s]+", "_", name)
    return name[:60] or "generated_document"


def set_document_styles(doc: Document):
    styles = doc.styles

    normal_style = styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    for style_name in ["Title", "Heading 1", "Heading 2"]:
        if style_name in styles:
            styles[style_name].font.name = "Calibri"

    if "Title" in styles:
        styles["Title"].font.size = Pt(20)

    if "Heading 1" in styles:
        styles["Heading 1"].font.size = Pt(14)

    if "Heading 2" in styles:
        styles["Heading 2"].font.size = Pt(12)


def add_bullet_list(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def create_docx(plan: dict) -> str:
    os.makedirs("output", exist_ok=True)

    title = plan.get("title", "Generated Business Document")
    document_type = plan.get("document_type", "business_report")
    assumptions = plan.get("assumptions", [])
    tasks = plan.get("tasks", [])
    sections = plan.get("sections", {})

    filename = f"{safe_filename(title)}.docx"
    filepath = os.path.join("output", filename)

    doc = Document()
    set_document_styles(doc)

    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Autonomous AI Document Agent Output"

    footer = section.footer.paragraphs[0]
    footer.text = "Email : shubhamsapkal2912@gmail.com"

    title_para = doc.add_paragraph(style="Title")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.add_run(title)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Document Type: {document_type.replace('_', ' ').title()}")

    doc.add_paragraph("")

    if assumptions:
        doc.add_heading("Assumptions", level=1)
        add_bullet_list(doc, assumptions)

    if tasks:
        doc.add_heading("Execution Plan", level=1)
        add_bullet_list(doc, tasks)

    for section_name, content in sections.items():
        doc.add_heading(section_name.replace("_", " ").title(), level=1)

        if isinstance(content, list):
            add_bullet_list(doc, content)
        else:
            doc.add_paragraph(str(content))

    doc.save(filepath)
    return filepath